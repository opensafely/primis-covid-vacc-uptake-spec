import csv
import json
import sys
from itertools import groupby
from docx import Document
from docx.table import Table

doc = Document(sys.argv[1])

rows = []

for table in doc.tables:
    if table.rows[0].cells[0].text not in ["Field Number", "147"]:
        continue

    for table_row in table.rows:
        if not table_row.cells[0].text:
            continue
        if table_row.cells[0].text == "Field Number":
            continue
        if table_row.cells[0].text.startswith("Note"):
            continue

        row = [
            " ".join(
                "".join(e.text for e in para._element.xpath(".//w:r"))
                for para in cell.paragraphs
            )
            .strip()
            .replace("–", "-")
            for cell in table_row.cells
        ]

        row = [item for item, _ in groupby(row)]
        rows.append(row)

extraction_criteria_records = []

last_id = 0

for id, group in groupby(rows, lambda row: row[0]):
    # print(id)
    id = int(id)
    assert id == last_id + 1
    last_id = id

    group = list(group)
    name = group[0][1].strip().upper()
    if not (name.endswith("_COD") or name in ["BMI_STAGE", "SEV_OBESITY"]):
        continue

    for row in group:
        assert row[1].strip().upper() == name

    record = {
        "id": id,
        "name": name,
        "details": group[0][2],
        "criteria": group[0][3],
    }

    if len(group) == 1:
        assert name == "BMI_COD"
        title = "BMI"
    else:
        title = group[1][2]
        if title[0] == "(":
            assert title[-1] == ")"
            title = title[1:-1]
    record["title"] = title

    extraction_criteria_records.append(record)


banding_and_grouping_records = []

last_banding_id = 0

for table in doc.tables:
    row0 = table.rows[0]

    if row0.cells[0].text.strip() == "1" and row0.cells[1].text.strip() == "Age Bands":
        inner_table = Table(table._element.xpath(".//w:tbl")[0], table._parent)
        record = {
            "banding_id": 1,
            "banding_name": "Age",
            "row": [
                [cell.text.strip() for cell in row.cells] for row in inner_table.rows
            ],
        }

        banding_and_grouping_records.append(record)
        continue

    if row0.cells[0].text.strip() == "":
        record = {
            "banding_id": int(table.rows[1].cells[0].text.strip()),
            "banding_name": row0.cells[1].text.strip(),
        }

        if record["banding_name"] == "Gender Bands":
            offset = 1
        else:
            offset = 2
            record["banding_key"] = table.rows[1].cells[1].text.strip()

        record["rows"] = [
            [cell.text.strip().replace("\n", "") for cell in row.cells[offset:]]
            for row in table.rows[1:]
        ]

        banding_and_grouping_records.append(record)
        continue

    try:
        banding_id = int(row0.cells[0].text.strip())
    except ValueError:
        continue

    if banding_id == 1:
        continue

    if banding_id == 3:
        record = {
            "banding_id": banding_id,
            "banding_name": row0.cells[1].text.strip(),
            "rows": [
                [cell.text.strip().replace("\n", "") for cell in row.cells]
                for row in table.rows[2:]
            ],
        }
        banding_and_grouping_records.append(record)
        continue

    if banding_id == 147:
        continue

    if banding_id < last_banding_id:
        break

    record = {
        "banding_id": banding_id,
        "banding_name": row0.cells[1].text.strip(),
        "banding_key": table.rows[1].cells[1].text.strip(),
        "rows": [
            [cell.text.strip().replace("\n", "") for cell in row.cells[2:]]
            for row in table.rows[1:]
        ],
    }
    banding_and_grouping_records.append(record)

    last_banding_id = banding_id


records = {
    "extraction_criteria": extraction_criteria_records,
    "bandings_and_groupings": banding_and_grouping_records,
}

with open(sys.argv[2], "w") as f:
    json.dump(records, f, indent=2)
